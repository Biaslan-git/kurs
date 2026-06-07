#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
                              QLineEdit, QTextEdit, QMessageBox, QComboBox, QFileDialog, QFrame)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CertificateWidget(QWidget):
    def __init__(self, db, go_back):
        super().__init__()
        self.db = db
        self.go_back = go_back
        self.selected_address_id = None
        self.build()

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter:
            self.generate_preview()
        else:
            super().keyPressEvent(event)
    
    def build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(15)
        
        # Заголовок
        header = QHBoxLayout()
        back_btn = QPushButton("← Назад")
        back_btn.setStyleSheet("QPushButton { background: transparent; color: #0071E3; border: none; font-size: 15px; padding: 5px 10px; } QPushButton:hover { background: #E5E5EA; border-radius: 6px; }")
        back_btn.setFixedWidth(100)
        back_btn.clicked.connect(self.go_back)
        header.addWidget(back_btn)
        
        title = QLabel("📄 Формирование справки")
        title.setFont(QFont("-apple-system", 22, QFont.Weight.Bold))
        header.addWidget(title)
        header.addStretch()
        layout.addLayout(header)
        
        hint = QLabel("✅ Последний шаг! Выберите адрес → Предпросмотр → Сохранить в Word")
        hint.setStyleSheet("color: #1D7A3A; background: #D4EDDA; border-radius: 8px; padding: 10px; font-size: 13px;")
        layout.addWidget(hint)
        
        # Параметры
        params_card = QFrame()
        params_card.setStyleSheet("background: white; border-radius: 12px; border: 1px solid #E5E5EA;")
        params_layout = QVBoxLayout(params_card)
        params_layout.setContentsMargins(15, 15, 15, 15)
        params_layout.setSpacing(10)
        
        row1 = QHBoxLayout()
        addr_lbl = QLabel("Адрес:")
        addr_lbl.setFixedWidth(80)
        addr_lbl.setStyleSheet("background: transparent; border: none; font-weight: 600;")
        row1.addWidget(addr_lbl)
        self.address_combo = QComboBox()
        row1.addWidget(self.address_combo, stretch=1)
        params_layout.addLayout(row1)
        
        row2 = QHBoxLayout()
        issued_lbl = QLabel("Выдал:")
        issued_lbl.setFixedWidth(80)
        issued_lbl.setStyleSheet("background: transparent; border: none; font-weight: 600;")
        row2.addWidget(issued_lbl)
        self.issued_by = QLineEdit()
        self.issued_by.setText("Специалист паспортного стола")
        row2.addWidget(self.issued_by)
        
        recip_lbl = QLabel("Получатель:")
        recip_lbl.setStyleSheet("background: transparent; border: none; font-weight: 600;")
        row2.addWidget(recip_lbl)
        self.recipient = QLineEdit()
        self.recipient.setPlaceholderText("ФИО получателя")
        row2.addWidget(self.recipient)
        params_layout.addLayout(row2)
        
        layout.addWidget(params_card)
        
        # Предпросмотр
        preview_title = QLabel("Предпросмотр:")
        preview_title.setFont(QFont("-apple-system", 13, QFont.Weight.Bold))
        layout.addWidget(preview_title)
        
        self.preview = QTextEdit()
        self.preview.setReadOnly(True)
        self.preview.setFont(QFont("Courier New", 11))
        self.preview.setStyleSheet("background: white; border: 1px solid #E5E5EA; border-radius: 10px; padding: 10px; color: #1D1D1F;")
        layout.addWidget(self.preview, stretch=2)
        
        # Кнопки
        btn_row = QHBoxLayout()
        
        preview_btn = QPushButton("👁️ Предпросмотр")
        preview_btn.setStyleSheet("QPushButton { background: #0071E3; border-radius: 10px; padding: 10px 20px; font-weight: 600; color: white; border: none; } QPushButton:hover { background: #0077ED; }")
        preview_btn.clicked.connect(self.generate_preview)
        
        save_btn = QPushButton("💾 Сохранить в Word")
        save_btn.setStyleSheet("QPushButton { background: #34C759; border-radius: 10px; padding: 10px 20px; font-weight: 600; color: white; border: none; } QPushButton:hover { background: #3DD768; }")
        save_btn.clicked.connect(self.save_word)
        
        btn_row.addWidget(preview_btn)
        btn_row.addWidget(save_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)
    
    def refresh(self):
        self.address_combo.clear()
        self.address_data = {}
        addresses = self.db.get_all_addresses()
        for addr in addresses:
            full = self.db.format_address(addr)
            self.address_combo.addItem(full)
            self.address_data[full] = addr['id']
    
    def generate_preview(self):
        text = self.address_combo.currentText()
        if text not in self.address_data:
            QMessageBox.warning(self, "Внимание", "Выберите адрес")
            return
        
        addr_id = self.address_data[text]
        address = self.db.get_address_by_id(addr_id)
        regs = self.db.get_active_registrations_by_address(addr_id)
        
        if not regs:
            QMessageBox.warning(self, "Внимание", "По этому адресу нет зарегистрированных лиц")
            return
        
        full_address = self.db.format_address(address)
        
        lines = []
        lines.append("=" * 60)
        lines.append("СПРАВКА О СОСТАВЕ СЕМЬИ".center(60))
        lines.append("(форма № 9)".center(60))
        lines.append("=" * 60)
        lines.append("")
        lines.append(f"Дата выдачи: {datetime.now().strftime('%d.%m.%Y')}")
        lines.append(f"Адрес: {full_address}")
        lines.append("")
        lines.append("Зарегистрированные лица:")
        lines.append("")
        
        for i, reg in enumerate(regs, 1):
            fio = f"{reg['last_name']} {reg['first_name']}"
            if reg['patronymic']:
                fio += f" {reg['patronymic']}"
            lines.append(f"{i}. {fio}")
            lines.append(f"   Дата рождения: {reg['birth_date']}")
            lines.append(f"   Паспорт: {reg['passport_series']} {reg['passport_number']}")
            lines.append(f"   Выдан: {reg['passport_issued_by']}, {reg['passport_issue_date']}")
            lines.append(f"   Зарегистрирован: {reg['registration_date']}")
            lines.append(f"   Родство: {reg['relationship']}")
            lines.append("")
        
        lines.append("=" * 60)
        lines.append(f"Выдал: {self.issued_by.text()}")
        lines.append("Подпись: _______________ М.П.")
        
        self.preview.setPlainText('\n'.join(lines))
        self.selected_address_id = addr_id
    
    def save_word(self):
        if not self.selected_address_id:
            QMessageBox.warning(self, "Внимание", "Сначала нажмите «Предпросмотр»")
            return
        if not self.recipient.text():
            QMessageBox.warning(self, "Внимание", "Укажите получателя")
            return
        
        path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить справку",
            f"Справка_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "Word документы (*.docx)"
        )
        
        if not path:
            return
        
        try:
            address = self.db.get_address_by_id(self.selected_address_id)
            regs = self.db.get_active_registrations_by_address(self.selected_address_id)
            
            doc = Document()
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(3)
                section.right_margin = Cm(1.5)
            
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = h.add_run("СПРАВКА О СОСТАВЕ СЕМЬИ\n(форма № 9)")
            r.bold = True
            r.font.size = Pt(14)
            
            doc.add_paragraph()
            doc.add_paragraph(f"Дата выдачи: {datetime.now().strftime('%d.%m.%Y')}")
            doc.add_paragraph(f"Адрес: {self.db.format_address(address)}")
            doc.add_paragraph()
            
            p = doc.add_paragraph("Зарегистрированные лица:")
            p.runs[0].bold = True
            doc.add_paragraph()
            
            for i, reg in enumerate(regs, 1):
                fio = f"{reg['last_name']} {reg['first_name']}"
                if reg['patronymic']:
                    fio += f" {reg['patronymic']}"
                p = doc.add_paragraph()
                p.add_run(f"{i}. {fio}\n").bold = True
                p.add_run(f"Дата рождения: {reg['birth_date']}\n")
                p.add_run(f"Паспорт: {reg['passport_series']} {reg['passport_number']}\n")
                p.add_run(f"Выдан: {reg['passport_issued_by']}, {reg['passport_issue_date']}\n")
                p.add_run(f"Зарегистрирован: {reg['registration_date']}\n")
                p.add_run(f"Родство: {reg['relationship']}")
                doc.add_paragraph()
            
            doc.add_paragraph()
            doc.add_paragraph(f"Выдал: {self.issued_by.text()}")
            doc.add_paragraph("Подпись: _______________ М.П.")
            
            doc.save(path)
            
            self.db.add_certificate(
                self.selected_address_id,
                datetime.now().strftime('%Y-%m-%d'),
                self.issued_by.text(),
                self.recipient.text()
            )
            
            QMessageBox.information(self, "✅ Готово!", f"Справка сохранена:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))
